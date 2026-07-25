import base64
import contextlib
import os
import tempfile
import threading
import unittest
import zipfile
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from unittest import mock

from app.retrieval import extract_images, html_to_readable
from app.tools import read_file, web_read


PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAIAAAACCAIAAAD91JpzAAAAFElEQVR4nGP8z8DAwMDAxMDAwMDAAAANHQEDasKb6QAAAABJRU5ErkJggg=="
)


class _Handler(BaseHTTPRequestHandler):
    pdf_bytes = b""

    def do_GET(self):
        if self.path == "/thread":
            body = """<!doctype html><html><body><nav>navigation noise</nav>
            <main class=thread><h1>Forum topic</h1><article>First post details</article>
            <article>Useful reply with Chinese evidence: 中文证据</article>
            <a href='/paper.pdf'>linked paper</a>
            <figure><img data-src='/diagram.png' alt='benchmark chart'>
            <figcaption>Figure 1 benchmark results</figcaption></figure></main></body></html>""".encode("utf-8")
            self._send("text/html", body)
        elif self.path == "/diagram.png":
            self._send("image/png", PNG_BYTES)
        elif self.path == "/paper.pdf":
            self._send("application/pdf", self.pdf_bytes)
        else:
            self.send_error(404)

    def _send(self, content_type, body):
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def log_message(self, *_args):
        pass


@contextlib.contextmanager
def serve(pdf_bytes=b""):
    _Handler.pdf_bytes = pdf_bytes
    server = ThreadingHTTPServer(("127.0.0.1", 0), _Handler)
    thread = threading.Thread(target=server.serve_forever, daemon=True)
    thread.start()
    try:
        yield f"http://127.0.0.1:{server.server_port}"
    finally:
        server.shutdown()
        server.server_close()
        thread.join(timeout=2)


class RetrievalTests(unittest.TestCase):
    def test_html_reading_preserves_forum_content_and_image_context(self):
        html = """<html><body><nav>noise</nav><aside>unrelated sidebar</aside><main><h1>Topic</h1>
        <p>First post</p><p>Second reply</p><a href='/paper.pdf'>Full paper</a>
        <figure><img src='plot.png' alt='plot alt'>
        <figcaption>Measured throughput</figcaption></figure></main></body></html>"""
        result = html_to_readable(html, "https://example.test/forum/thread")
        self.assertIn("First post\nSecond reply", result)
        self.assertIn("https://example.test/forum/plot.png", result)
        self.assertIn("Measured throughput", result)
        self.assertIn("Full paper - https://example.test/paper.pdf", result)
        self.assertNotIn("noise", result)
        self.assertNotIn("unrelated sidebar", result)

    def test_web_read_and_extract_lazy_loaded_forum_image(self):
        with tempfile.TemporaryDirectory() as tmp, serve() as base:
            text = web_read(f"{base}/thread")
            self.assertIn("Useful reply with Chinese evidence", text)
            self.assertIn("中文证据", text)
            self.assertIn(f"{base}/diagram.png", text)
            self.assertIn(f"linked paper - {base}/paper.pdf", text)

            result = extract_images(f"{base}/thread", Path(tmp), max_images=2)
            self.assertIn("benchmark results", result)
            files = list(Path(tmp).glob("*.png"))
            self.assertEqual(1, len(files))
            self.assertEqual(PNG_BYTES, files[0].read_bytes())

    def test_remote_pdf_is_read_with_page_markers_and_cached(self):
        try:
            import fitz
        except ImportError:
            self.skipTest("PyMuPDF is not installed")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Remote paper page one")
        pdf_bytes = doc.tobytes()
        doc.close()

        with tempfile.TemporaryDirectory() as tmp, serve(pdf_bytes) as base:
            with mock.patch.dict(os.environ, {"APPDATA": tmp}):
                result = web_read(f"{base}/paper.pdf")
        self.assertIn("远程 PDF 已缓存", result)
        self.assertIn("第 1 页", result)
        self.assertIn("Remote paper page one", result)

    def test_pdf_pages_are_rendered_for_vision(self):
        try:
            import fitz
        except ImportError:
            self.skipTest("PyMuPDF is not installed")
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            pdf_path = tmp_path / "paper.pdf"
            doc = fitz.open()
            for index in range(2):
                page = doc.new_page()
                page.insert_text((72, 72), f"Chart page {index + 1}")
            doc.save(pdf_path)
            doc.close()

            result = extract_images(str(pdf_path), tmp_path / "out", pages="2", max_images=2)
            self.assertIn("第 2 页", result)
            self.assertNotIn("第 1 页", result)
            self.assertEqual(1, len(list((tmp_path / "out").glob("*.png"))))

            text = read_file(str(pdf_path), pages="2")
            self.assertIn("Chart page 2", text)
            self.assertNotIn("Chart page 1", text)

    def test_docx_embedded_media_is_extracted(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            docx_path = tmp_path / "sample.docx"
            with zipfile.ZipFile(docx_path, "w") as archive:
                archive.writestr("[Content_Types].xml", "<Types/>")
                archive.writestr("word/document.xml", "<document/>")
                archive.writestr("word/media/image1.png", PNG_BYTES)

            result = extract_images(str(docx_path), tmp_path / "out")
            self.assertIn("Word 内嵌图片 1", result)
            extracted = list((tmp_path / "out").glob("*.png"))
            self.assertEqual(1, len(extracted))
            self.assertEqual(PNG_BYTES, extracted[0].read_bytes())

    def test_extraction_can_run_existing_local_ocr_in_same_call(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            image_path = tmp_path / "screenshot.png"
            image_path.write_bytes(PNG_BYTES)
            with mock.patch("app.vision.ocr_image", return_value="Detected local text") as mocked:
                result = extract_images(
                    str(image_path), tmp_path / "out", use_ocr=True,
                )

            mocked.assert_called_once()
            self.assertIn("Detected local text", result)
            self.assertIn("优先根据正文、图注和 OCR 文字回答", result)

    def test_docx_tables_are_searchable_text(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "table.docx"
            doc = Document()
            doc.add_paragraph("Report summary")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Metric"
            table.cell(0, 1).text = "42"
            doc.save(path)

            result = read_file(str(path))
            self.assertIn("Report summary", result)
            self.assertIn("=== 表格 1 ===", result)
            self.assertIn("Metric\t42", result)


if __name__ == "__main__":
    unittest.main()
